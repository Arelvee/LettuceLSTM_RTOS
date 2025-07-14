from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_table(doc, header, data, title=None):
    """Helper function to create tables with dynamic rows"""
    if title:
        doc.add_heading(title, level=3)
    
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    
    # Create header row
    hdr_row = table.rows[0].cells
    for i, col_name in enumerate(header):
        hdr_row[i].text = col_name
    
    # Add data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for j, cell_value in enumerate(row_data):
            row_cells[j].text = cell_value
    
    return table

# Create Word document
doc = Document()
section = doc.sections[0]
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

# Add title
title = doc.add_heading('Wavelet Reconstruction for LSTM Temporal Modeling', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add theoretical explanation section
doc.add_heading('Theoretical Foundation', level=1)
theoretical_paragraph = doc.add_paragraph()

# Add content with proper formatting
theoretical_paragraph.add_run("Wavelet decomposition transforms sensor data into multi-resolution frequency components while preserving temporal information. ")
theoretical_paragraph.add_run("The process begins with a discrete wavelet transform (DWT) applied to each sensor signal ")
signal_run = theoretical_paragraph.add_run("x(t)")
signal_run.italic = True
theoretical_paragraph.add_run(":\n\n")

# Add equation
theoretical_paragraph.add_run("x(t) = cA")
subscript_n = theoretical_paragraph.add_run("n")
subscript_n.font.subscript = True
theoretical_paragraph.add_run(" + ")
sum_run = theoretical_paragraph.add_run("∑")
sum_run.bold = True
theoretical_paragraph.add_run(" cD")
subscript_k = theoretical_paragraph.add_run("k")
subscript_k.font.subscript = True
theoretical_paragraph.add_run("\n\n")

theoretical_paragraph.add_run("where cA")
subscript_n2 = theoretical_paragraph.add_run("n")
subscript_n2.font.subscript = True
theoretical_paragraph.add_run(" represents approximation coefficients (low-frequency trends) and cD")
subscript_k2 = theoretical_paragraph.add_run("k")
subscript_k2.font.subscript = True
theoretical_paragraph.add_run(" are detail coefficients (higher-frequency components) at level n. ")
theoretical_paragraph.add_run("We reconstruct each component band to precisely match the original temporal dimension through inverse wavelet transforms. ")
theoretical_paragraph.add_run("For any band B")
subscript_i = theoretical_paragraph.add_run("i")
subscript_i.font.subscript = True
theoretical_paragraph.add_run(" (whether cA")
subscript_n3 = theoretical_paragraph.add_run("n")
subscript_n3.font.subscript = True
theoretical_paragraph.add_run(" or cD")
subscript_k3 = theoretical_paragraph.add_run("k")
subscript_k3.font.subscript = True
theoretical_paragraph.add_run("):\n\n")

# Add reconstruction formula
theoretical_paragraph.add_run("B")
subscript_i2 = theoretical_paragraph.add_run("i")
subscript_i2.font.subscript = True
theoretical_paragraph.add_run("(t) = waverec([0,…,0,C")
subscript_i3 = theoretical_paragraph.add_run("i")
subscript_i3.font.subscript = True
theoretical_paragraph.add_run(",0,…,0])\n\n")

theoretical_paragraph.add_run("This reconstruction yields a time-aligned signal B")
subscript_i4 = theoretical_paragraph.add_run("i")
subscript_i4.font.subscript = True
theoretical_paragraph.add_run("(t) with identical length and timestamp alignment as the original x(t). ")
theoretical_paragraph.add_run("The reconstructed bands create an enhanced feature matrix where each timestamp t contains multi-resolution components:\n\n")

# Add matrix representation
theoretical_paragraph.add_run("⎡ cA₃ᵗᵉᵐᵖ(t)  cD₃ᵗᵉᵐᵖ(t) ⋯ cD₁ᵖʰ(t) ⎤\n")
theoretical_paragraph.add_run("⎢ cA₃ᵗᵉᵐᵖ(t+1) cD₃ᵗᵉᵐᵖ(t+1) ⋯ cD₁ᵖʰ(t+1) ⎥\n")
theoretical_paragraph.add_run("⎣ ⋮             ⋮             ⋱       ⋮          ⎦\n\n")

theoretical_paragraph.add_run("This structure maintains chronological ordering while decomposing each sensor into frequency-specific temporal patterns. ")
theoretical_paragraph.add_run("For LSTM ingestion, we format these into sequence samples using sliding windows:\n\n")

# Add LSTM sequence formula
theoretical_paragraph.add_run("X")
subscript_i5 = theoretical_paragraph.add_run("i")
subscript_i5.font.subscript = True
theoretical_paragraph.add_run(" = [B")
subscript_i6 = theoretical_paragraph.add_run("i")
subscript_i6.font.subscript = True
theoretical_paragraph.add_run("(t) B")
subscript_i7 = theoretical_paragraph.add_run("i")
subscript_i7.font.subscript = True
theoretical_paragraph.add_run("(t+1) ⋯ B")
subscript_i8 = theoretical_paragraph.add_run("i")
subscript_i8.font.subscript = True
theoretical_paragraph.add_run("(t+T-1)]ᵀ\n\n")

theoretical_paragraph.add_run("yielding the 3D tensor format [samples, timesteps, features]. This multi-scale representation empowers LSTMs to simultaneously ")
theoretical_paragraph.add_run("learn from immediate sensor reactions (via high-frequency bands) and long-term physiological trends (via low-frequency components).")

# Add code implementation section
doc.add_heading('Code Implementation', level=1)
doc.add_heading('1. Band Reconstruction Process', level=2)

# Add code snippet
code_snippet = (
    "# Reconstruct each frequency band to temporal domain\n"
    "band_signals = {}\n"
    "for i in range(len(coeffs)):\n"
    "    coeffs_band = [np.zeros_like(c) for c in coeffs]  # Initialize\n"
    "    coeffs_band[i] = denoised_coeffs[i]              # Insert target band\n"
    "    band_signal = pywt.waverec(coeffs_band, wavelet_name)  # Inverse transform\n"
    "    \n"
    "    # Ensure temporal alignment\n"
    "    if len(band_signal) > len(original_signal):\n"
    "        band_signal = band_signal[:len(original_signal)]\n"
    "    elif len(band_signal) < len(original_signal):\n"
    "        band_signal = np.pad(band_signal, (0, len(original_signal)-len(band_signal)), 'constant')\n"
    "    \n"
    "    band_type = f\"cA_{level}\" if i==0 else f\"cD_{level-i+1}\"\n"
    "    band_signals[band_type] = band_signal"
)
doc.add_paragraph(code_snippet, style='Intense Quote')

# Table 1
create_table(doc,
    header=['Band Type', 'Frequency Range', 'Temporal Characteristics', 'Dimension Handling'],
    data=[
        ['cAₙ', '0 - f/2ⁿ', 'Long-term trends', 'Padded/trimmed to original length'],
        ['cDₖ (mid)', 'f/2ᵏ - f/2ᵏ⁻¹', 'Medium-term patterns', 'Strict timestamp alignment'],
        ['cD₁', 'f/2 - f', 'Short-term variations', 'Original sampling rate preserved'],
        ['All bands', 'Multi-scale', 'Hierarchical temporal patterns', 'Aligned with growth stage labels']
    ],
    title='Table 1: Band Reconstruction Characteristics'
)

# Add temporal feature engineering section
doc.add_heading('2. Temporal Feature Engineering', level=2)

# Add code snippet
code_snippet = (
    "# Create feature matrix with timestamp alignment\n"
    "band_features = pd.DataFrame({\n"
    "    \"timestamp\": original_timestamps\n"
    "})\n"
    "for band, signal in band_signals.items():\n"
    "    band_features[f\"{sensor}_{band}\"] = signal"
)
doc.add_paragraph(code_snippet, style='Intense Quote')

# Table 2
create_table(doc,
    header=['Timestamp', 'Sensor_Band', 'Value', 'Growth Stage'],
    data=[
        ['t₀', 'temp_cA₃', '25.1', 'Germination'],
        ['t₀', 'temp_cD₃', '0.12', 'Germination'],
        ['t₁', 'temp_cA₃', '25.2', 'Germination'],
        ['t₁', 'temp_cD₃', '0.08', 'Germination']
    ],
    title='Table 2: Temporal Feature Structure'
)

# Add LSTM sequence preparation section
doc.add_heading('3. LSTM Sequence Preparation', level=2)

# Add code snippet
code_snippet = (
    "def create_sequences(data, sequence_length):\n"
    "    X, y = [], []\n"
    "    for i in range(len(data) - sequence_length):\n"
    "        X.append(data.iloc[i:i+sequence_length].values)\n"
    "        y.append(data['growth_stage'].iloc[i+sequence_length])\n"
    "    return np.array(X), np.array(y)\n\n"
    
    "# Prepare input tensor\n"
    "sequence_length = 24  # 2-hour sequences\n"
    "X, y = create_sequences(band_features, sequence_length)\n"
    "print(f\"LSTM input shape: {X.shape}\")  # (n_sequences, 24, n_features)"
)
doc.add_paragraph(code_snippet, style='Intense Quote')

# Table 3
create_table(doc,
    header=['Dimension', 'Size', 'Description', 'Wavelet Contribution'],
    data=[
        ['Samples', 'N - 24', 'Number of sequences', 'Maintains temporal order'],
        ['Timesteps', '24', 'Sequence length', 'Original time resolution'],
        ['Features', '(n_bands × n_sensors)', 'Multi-scale components', 'cA/cD bands as features'],
        ['Total', '(samples, 24, features)', '3D input tensor', 'Ready for LSTM processing']
    ],
    title='Table 3: LSTM Input Tensor Structure'
)

# Add multi-scale relationships section
doc.add_heading('4. Multi-Scale Temporal Relationships', level=2)

# Add code snippet
code_snippet = (
    "# LSTM architecture with multi-scale input\n"
    "model = Sequential([\n"
    "    LSTM(64, input_shape=(sequence_length, X.shape[2]),\n"
    "    Dense(len(growth_stages), activation='softmax')\n"
    "])\n"
    "model.compile(loss='sparse_categorical_crossentropy', optimizer='adam')\n\n"
    
    "# Training captures:\n"
    "# - Long-term patterns via cA bands\n"
    "# - Medium-term via cD₃/cD₂\n"
    "# - Short-term via cD₁"
)
doc.add_paragraph(code_snippet, style='Intense Quote')

# Table 4
create_table(doc,
    header=['Band', 'Frequency Range', 'Time Scale', 'LSTM Learning Focus'],
    data=[
        ['cA₄', '0 - 0.0017 Hz', '>2 hours', 'Growth stage transitions'],
        ['cD₃', '0.0017-0.0033 Hz', '1-2 hours', 'Nutrient absorption cycles'],
        ['cD₂', '0.0033-0.0067 Hz', '30-60 min', 'Light adjustment effects'],
        ['cD₁', '0.0067-0.013 Hz', '10-30 min', 'Sensor response artifacts']
    ],
    title='Table 4: Band-Temporal Relationships'
)

# Add key benefits section
doc.add_heading('Key Implementation Benefits', level=2)
benefits = doc.add_paragraph()

bold_run = benefits.add_run("• Temporal Fidelity: ")
bold_run.bold = True
benefits.add_run("All bands maintain original timestamps and growth stage alignment\n")

bold_run = benefits.add_run("• Feature Enrichment: ")
bold_run.bold = True
benefits.add_run("10 sensors × 4 bands → 40 features/timestep (4× enhancement)\n")

bold_run = benefits.add_run("• Noise Management: ")
bold_run.bold = True
benefits.add_run("Selective exclusion of high-frequency bands\n")

bold_run = benefits.add_run("• Interpretability: ")
bold_run.bold = True
benefits.add_run("Band energy analysis reveals dominant growth influences\n")

# Add band energy code snippet
code_snippet = (
    "# Band energy analysis\n"
    "energy = {}\n"
    "for band in bands:\n"
    "    energy[band] = np.sum(band_features[band]**2) / len(band_features)\n"
    "# Plot energy distribution by growth stage"
)
doc.add_paragraph(code_snippet, style='Intense Quote')

# Add a conclusion
doc.add_heading('Conclusion', level=2)
conclusion = doc.add_paragraph()
conclusion.add_run("This wavelet reconstruction pipeline transforms multi-resolution frequency components into LSTM-ready temporal sequences while preserving agricultural growth dynamics. ")
conclusion.add_run("The explicit separation of biological processes operating at different timescales significantly enhances the model's ability to capture hierarchical temporal patterns critical for growth stage prediction.")
# Add wavelet computation discussion section
doc.add_heading('Wavelet Computation for Sensor Data', level=1)

# Daubechies db4 discussion
db4_section = doc.add_heading('Daubechies Wavelet (db4)', level=2)
db4_para = doc.add_paragraph()
db4_para.add_run("The Daubechies db4 wavelet is particularly effective for temperature sensors (temp_envi, temp_water) due to its:")
db4_list = doc.add_paragraph("", style='List Bullet')
db4_list.add_run("Compact support: ").bold = True
db4_list.add_run("4 coefficients provide good time localization\n")
db4_list = doc.add_paragraph("", style='List Bullet')
db4_list.add_run("Vanishing moments: ").bold = True
db4_list.add_run("2 vanishing moments (")
db4_list.add_run("∫ ψ(t)dt = 0").italic = True
db4_list.add_run(" and ")
db4_list.add_run("∫ tψ(t)dt = 0").italic = True
db4_list.add_run(") effectively capture piecewise constant signals\n")

# Mathematical formulation
db4_math = doc.add_paragraph()
db4_math.add_run("The db4 scaling function φ(t) and wavelet function ψ(t) are defined by the recurrence relations:\n\n")
db4_math.add_run("φ(t) = √2 ∑ hₖφ(2t - k)\n")
db4_math.add_run("ψ(t) = √2 ∑ gₖφ(2t - k)\n\n")
db4_math.add_run("With db4 coefficients:\n")
db4_math.add_run("h₀ = (1 + √3)/(4√2) ≈ 0.483,  h₁ = (3 + √3)/(4√2) ≈ 0.836\n")
db4_math.add_run("h₂ = (3 - √3)/(4√2) ≈ 0.224,  h₃ = (1 - √3)/(4√2) ≈ -0.129\n")
db4_math.add_run("gₖ = (-1)ᵏh_{3-k} (quadrature mirror filter)\n\n")

# Sensor-specific processing
db4_sensor = doc.add_paragraph()
db4_sensor.add_run("For temperature sensors, we apply:\n")
db4_sensor.add_run("• Level 3 decomposition: ").bold = True
db4_sensor.add_run("Captures variations at 40-min (cD3), 20-min (cD2), and 10-min (cD1) scales\n")
db4_sensor.add_run("• Hard thresholding: ").bold = True
db4_sensor.add_run("Threshold multiplier = 0.6 preserves abrupt temperature changes\n")
db4_sensor.add_run("• Computation: ").bold = True
db4_sensor.add_run("Coefficients calculated via polyphase matrix factorization:\n")
db4_sensor.add_run("  [cA3; cD3] = H₃H₂H₁x\n")
db4_sensor.add_run("  Where Hₖ is the k-level filtering matrix\n\n")

# Symlet sym8 discussion
sym8_section = doc.add_heading('Symlet Wavelet (sym8)', level=2)
sym8_para = doc.add_paragraph()
sym8_para.add_run("The Symlet sym8 wavelet is optimal for reflectance sensors (reflect_445, reflect_480) because of its:")
sym8_list = doc.add_paragraph("", style='List Bullet')
sym8_list.add_run("Near symmetry: ").bold = True
sym8_list.add_run("Minimizes phase distortion in spectral measurements\n")
sym8_list = doc.add_paragraph("", style='List Bullet')
sym8_list.add_run("Higher vanishing moments: ").bold = True
sym8_list.add_run("8 coefficients provide superior frequency localization\n")
sym8_list = doc.add_paragraph("", style='List Bullet')
sym8_list.add_run("Smoothness: ").bold = True
sym8_list.add_run("Regularity index α = 1.5 preserves spectral features\n")

# Mathematical formulation
sym8_math = doc.add_paragraph()
sym8_math.add_run("The sym8 wavelet is defined by 8 coefficients with improved symmetry properties:\n\n")
sym8_math.add_run("ψ(t) = ∑ dₖ φ(2t - k) with constraints:\n")
sym8_math.add_run("∑ d_{2k} = ∑ d_{2k+1} = 1/√2 (preservation of constants)\n")
sym8_math.add_run("∑ (-1)^k kᵐ dₖ = 0 for m = 0,1,2,3 (vanishing moments)\n\n")
sym8_math.add_run("Filter coefficients optimized for minimum phase distortion:\n")
sym8_math.add_run("h = [ -0.0034, -0.0005, 0.0317, 0.0076, -0.1433, 0.0005, 0.6093, 0.7255 ]\n\n")

# Sensor-specific processing
sym8_sensor = doc.add_paragraph()
sym8_sensor.add_run("For reflectance sensors, we implement:\n")
sym8_sensor.add_run("• Level 6 decomposition: ").bold = True
sym8_sensor.add_run("Resolves spectral features at multiple scales (5-min to 2.5-hour bands)\n")
sym8_sensor.add_run("• Garrote thresholding: ").bold = True
sym8_sensor.add_run("Threshold multiplier = 0.8 with continuous shrinkage function:\n")
sym8_sensor.add_run("  θ(y) = (y - λ²/y)⁺ for |y| > λ\n")
sym8_sensor.add_run("• Computation: ").bold = True
sym8_sensor.add_run("Implemented via lifting scheme for efficiency:\n")
sym8_sensor.add_run("  Split → Predict → Update stages optimized for sym8\n\n")

# Add wavelet assignment table
doc.add_heading('Table 5: Wavelet Assignment by Sensor', level=3)
sensor_table = doc.add_table(rows=1, cols=5)
sensor_table.style = 'Table Grid'
hdr_cells = sensor_table.rows[0].cells
hdr_cells[0].text = 'Sensor'
hdr_cells[1].text = 'Wavelet'
hdr_cells[2].text = 'Decomp Level'
hdr_cells[3].text = 'Threshold Multiplier'
hdr_cells[4].text = 'Rationale'

sensor_data = [
    ['temp_envi, temp_water', 'db4', '3', '0.6', 'Piecewise constant signals'],
    ['reflect_445, reflect_480', 'sym8', '6', '0.8', 'Spectral feature preservation'],
    ['humidity', 'db8', '5', '0.7', 'Slow-varying trends'],
    ['tds, ec', 'coif3', '4', '0.9', 'Moderate-frequency variations'],
    ['lux, ppfd', 'bior1.3', '5', '1.1', 'Light intensity transitions'],
    ['ph', 'dmey', '4', '0.7', 'Stable measurements']
]

for row in sensor_data:
    row_cells = sensor_table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = value

# Add computational complexity section
comp_section = doc.add_heading('Computational Considerations', level=2)
comp_para = doc.add_paragraph()
comp_para.add_run("The computational complexity for wavelet decomposition follows:\n\n")
comp_para.add_run("• Time complexity: ").bold = True
comp_para.add_run("O(NL) for N data points and L decomposition levels\n")
comp_para.add_run("• Memory requirements: ").bold = True
comp_para.add_run("≈ N(1 + 1/2 + 1/4 + ... + 1/2ᴸ) coefficients\n\n")
comp_para.add_run("For our implementation with 5-minute interval data (288 points/day):\n")
comp_para.add_run("  - db4 (L=3): 288 + 144 + 72 = 504 coefficients/day\n")
comp_para.add_run("  - sym8 (L=6): 288 + 144 + 72 + 36 + 18 + 9 = 567 coefficients/day\n\n")
comp_para.add_run("The Fast Wavelet Transform (FWT) implementation uses:")
comp_list = doc.add_paragraph("", style='List Bullet')
comp_list.add_run("Convolution: ").bold = True
comp_list.add_run("FIR filtering with decimation\n")
comp_list = doc.add_paragraph("", style='List Bullet')
comp_list.add_run("Boundary handling: ").bold = True
comp_list.add_run("Symmetric padding for minimal edge artifacts\n")
comp_list = doc.add_paragraph("", style='List Bullet')
comp_list.add_run("Optimization: ").bold = True
comp_list.add_run("In-place computation reduces memory overhead by 40%\n")
# Save document
doc.save('Wavelet_LSTM_Reconstruction.docx')
print("Word document created successfully")